# ===============================================================
# app.py - Your Klarity Architect Template Factory
# ===============================================================

import streamlit as st
import openai
from docx import Document
from io import BytesIO
import json
import sys
from importlib.metadata import version
from docx.shared import Pt, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn

# --- Helper Functions (The "Workers") ---

# This function creates the special prompt we send to OpenAI
def create_meta_prompt(master_context, input_files, section_title, section_goal, output_format):
    """Creates a detailed prompt for OpenAI to generate the section config."""
    
    # Build the input files context
    input_files_context = ""
    if input_files:
        input_files_context = "**Available Input Files:**\n"
        for input_file in input_files:
            input_files_context += f"- **{input_file['name']}** ({input_file['type']}): {input_file['description']}\n"
    
    return f"""
    You are an AI assistant specialized in creating configuration prompts for a tool called Klarity Architect.
    Your task is to generate the configuration for a single section of a template.

    **Master Template Context:**
    ---
    {master_context}
    ---

    {input_files_context}

    **Section Details to Configure:**
    *   **Section Title:** "{section_title}"
    *   **Plain Language Goal:** "{section_goal}"
    *   **Required Output Format:** "{output_format}"

    **Your Instructions:**
    Based on all the information above, you must generate a detailed, expert-level prompt that a second LLM will use to extract information from documents. This prompt MUST be self-contained and structured with **Role**, **Context**, **Task**, and **Instructions**. It should be highly specific and actionable.

    **CRITICAL: You MUST follow this EXACT template format. Do not deviate from the line breaks or structure:**

    TEMPLATE:
    "*Role:* [Your role description here]\\n\\n*Context:* [Your context description here]\\n\\n*Task:* [Your task description here]\\n\\n*Instructions:*\\n\\n1. [First instruction]\\n2. [Second instruction]\\n3. [Third instruction]"

    **REQUIREMENTS:**
    - Use *Role:*, *Context:*, *Task:*, and *Instructions:* (with asterisks)
    - You MUST use \\n\\n (double line breaks) between sections: after Role, after Context, after Task
    - You MUST use \\n\\n (double line breaks) after *Instructions:* and before the first numbered item
    - Use \\n (single line break) for bullet points within instructions
    - Number the instructions with 1., 2., 3., etc.
    - Use bullet points (*) for sub-items within instructions
    - Consider the available input files when crafting the Context and Instructions sections

    In addition to the main prompt, determine the best 'sub_type' for the '{output_format}'.
    *   If format is 'Text', the sub_type must be 'freeform' or 'bulleted'.
    *   If format is 'Table', the sub_type must be 'default'.

    **CRITICAL OUTPUT FORMAT:**
    You must respond with ONLY a single, valid JSON object. Do not add any conversational text, explanations, or markdown formatting like ```json. The JSON object must have these exact keys: "type", "sub_type", "prompt", "include_screenshots", "screenshot_instructions".

    Example of a perfect response:
    {{
        "type": "text",
        "sub_type": "freeform",
        "prompt": "*Role:* You are an expert Senior Implementation Consultant at Tekion, specializing in large, complex dealer groups. You are highly proficient in both legacy DMS structures (like CDK and DealerTrack) and the advanced capabilities of Tekion's Automotive Retail Cloud.\\n\\n*Context:* You are reviewing the provided discovery call transcript(s) with a new dealership prospect. You have already been provided with a comprehensive *Tekion Knowledge Base* in the 'Additional Instructions', detailing all Tekion platform features, terminology (Instances, Sites, Prefixes, BOC), and standard configurations. You must leverage this internal knowledge to interpret the dealer's current state and its implications for a Tekion implementation.\\n\\n*Task:* Your task is to draft the *Executive Summary & Key Implementation Themes* section for an internal implementation blueprint document. This section must be concise and strategic, designed for a Tekion Engagement Manager who needs to quickly understand the core nature of the project.\\n\\n*Instructions:*\\n\\n1. *Executive Summary:* First, write a brief paragraph summarizing the dealership's profile. Include:\\n * The dealer group's name.\\n * Their scale (e.g., number of rooftops, brands mentioned).\\n * Their current DMS (e.g., CDK, Autosoft).\\n * Their core operational and accounting structure (e.g., centralized accounting under a BOC, single tax ID, how they currently structure multi-brand companies).\\n\\n2. *Key Implementation Themes:* After the summary, identify and articulate the *3-4 most critical, overarching themes* for this implementation. These are not just individual findings but the strategic pillars that will define the project's complexity and success. For each theme, provide a brief sentence explaining its significance.\\n\\n * *Focus on Synthesis, Not Just Summary:* Do not list every detail. Synthesize related issues into a single strategic theme.\\n * *Examples of good themes:*\\n * `Complex Company Restructuring:` if the dealer has many brands combined under single legacy company numbers that need to be untangled into proper Tekion Sites or Instances.\\n * `Significant Process Modernization:` if the dealer relies heavily on manual workarounds, Excel, or outdated third-party tools (like Filebound) that will be replaced by Tekion.\\n * `Navigating Non-Standard Accounting Workflows:` if the dealer has unique intercompany transaction methods or treats internal entities (like a body shop) as external vendors.\\n * `Critical State-Specific Configuration:` if there are unique tax, title, or legal requirements (like the Pennsylvania trade-in tax issue) that demand precise setup.",
        "include_screenshots": "no",
        "screenshot_instructions": "none"
    }}
    """

# This function formats the comment string exactly as Klarity Architect needs it
def format_comment_string(json_string):
    """Takes the JSON response from OpenAI and formats it into the Klarity comment spec."""
    try:
        data = json.loads(json_string)
        # Get the prompt content and ensure it uses the correct formatting
        prompt_content = data.get('prompt', '')
        
        # Convert markdown bold (**text**) to asterisk format (*text*) for consistency
        prompt_content = prompt_content.replace('**', '*')
        
        # CRITICAL FIX: Convert actual newlines back to escaped \n characters
        # json.loads() converts \n to actual newlines, but we need literal \n in the comment
        prompt_content = prompt_content.replace('\r\n', '\n')   # normalize first
        prompt_content = prompt_content.replace('\n', '\\n')    # convert to escaped \n
        
        # The prompt content should be exactly as it comes from the AI, with proper line breaks
        comment_text = f"""type - {data.get('type', '')}
sub_type - {data.get('sub_type', '')}
prompt - {prompt_content}
include_screenshots - {data.get('include_screenshots', 'no')}
screenshot_instructions - {data.get('screenshot_instructions', 'none')}"""
        return comment_text
    except json.JSONDecodeError:
        st.error("AI returned an invalid format. Could not parse the comment data.")
        return None

# This function builds the final .docx file with corrected comment creation
def create_docx(template_title, processed_sections):
    """Builds the Word document with properly formatted comments."""
    document = Document()
    # Remove the template title header - we only want section headers
    
    for section in processed_sections:
        # Check if we have a comment string to add.
        if section.get('comment_string'):
            # Add the section title as a header
            heading = document.add_heading(section['title'], level=1)
            
            # Get the run from the heading to anchor the comment to
            run = heading.runs[0] if heading.runs else heading.add_run(section['title'])
            
            # Create the comment using the correct Document.add_comment() method
            # This method takes runs as the first parameter, then text, author, initials
            comment = document.add_comment(
                runs=[run],
                text=section['comment_string'],
                author="System",
                initials="S"
            )
        else:
            # If there's no comment, just add the title as a header.
            document.add_heading(section['title'], level=1)

    # Save the document to an in-memory stream.
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- Main Streamlit App UI and Logic ---

# Set the page title and a fun icon
st.set_page_config(page_title="Klarity Template Factory", page_icon="🏭")

# The App Title
st.title("🏭 Klarity Architect Template Factory")

# --- Step 1: Master Context ---
st.header("Step 1: Define the Master Context")
st.info("Provide the overall purpose for this template. This guides the AI for all sections.", icon="🎯")

master_context = st.text_area(
    "Master Context for the Template",
    height=150,
    placeholder="e.g., This template creates a Process Definition Document after a discovery call with a new automotive dealership client implementing the Tekion platform. The output is for an internal Engagement Manager at Tekion."
)

# --- Step 2: Input Files Context ---
st.header("Step 2: Define Input Files")
st.info("Add the input files that will be used to generate this template. This helps the AI understand what data sources are available.", icon="📁")

# Initialize session state for input files
if 'input_files' not in st.session_state:
    st.session_state.input_files = [{
        'name': 'Discovery Call Transcript',
        'type': 'DOCX',
        'description': 'Meeting notes and transcript from the initial discovery call with the client, including their current system details and requirements.'
    }]

# Display the UI for each input file
for i, input_file in enumerate(st.session_state.input_files):
    cols = st.columns([3, 2, 4, 1])
    with cols[0]:
        st.session_state.input_files[i]['name'] = st.text_input("Input File Name", value=input_file['name'], key=f"input_name_{i}")
    with cols[1]:
        st.session_state.input_files[i]['type'] = st.selectbox("File Type", options=['PDF', 'DOCX', 'XLSX', 'MP4', 'TXT', 'CSV'], index=0, key=f"input_type_{i}")
    with cols[2]:
        st.session_state.input_files[i]['description'] = st.text_input("File Description", value=input_file['description'], key=f"input_desc_{i}")
    with cols[3]:
        if st.button("❌", key=f"del_input_{i}", help="Delete this input file"):
            st.session_state.input_files.pop(i)
            st.rerun()

if st.button("➕ Add Input File"):
    st.session_state.input_files.append({'name': '', 'type': 'DOCX', 'description': ''})
    st.rerun()

# --- Step 3: Define Sections ---
st.header("Step 3: Add and Define Your Sections")
st.info("Add a row for each section of your final document. Describe its goal in plain English.", icon="✍️")

# Initialize session state for our list of sections
if 'sections' not in st.session_state:
    st.session_state.sections = [{
        'title': 'Executive Summary & Key Themes',
        'format': 'Text (Freeform)',
        'goal': 'Summarize the client\'s profile (name, size, current system) and list the 3-4 most important, high-level implementation themes.'
    }]

# Display the UI for each section
for i, section in enumerate(st.session_state.sections):
    cols = st.columns([3, 2, 4, 1])
    with cols[0]:
        st.session_state.sections[i]['title'] = st.text_input("Section Title", value=section['title'], key=f"title_{i}")
    with cols[1]:
        st.session_state.sections[i]['format'] = st.selectbox("Output Format", options=['Text (Freeform)', 'Text (Bulleted)', 'Table'], index=0, key=f"format_{i}")
    with cols[2]:
        st.session_state.sections[i]['goal'] = st.text_input("Section Goal (Plain English)", value=section['goal'], key=f"goal_{i}")
    with cols[3]:
        if st.button("❌", key=f"del_{i}", help="Delete this section"):
            st.session_state.sections.pop(i)
            st.rerun()

if st.button("➕ Add Section"):
    st.session_state.sections.append({'title': '', 'format': 'Text (Freeform)', 'goal': ''})
    st.rerun()

st.divider()

# --- Step 4: Generate the Template ---
st.header("Step 4: Generate and Download")
if st.button("✨ Generate Template ✨", type="primary", use_container_width=True):
    # Enhanced validation
    if not master_context.strip():
        st.warning("Please provide a Master Context before generating.")
    elif not st.session_state.input_files:
        st.warning("Please add at least one input file.")
    elif any(not input_file['name'].strip() for input_file in st.session_state.input_files):
        st.warning("Please provide names for all input files.")
    elif any(not input_file['description'].strip() for input_file in st.session_state.input_files):
        st.warning("Please provide descriptions for all input files.")
    elif not st.session_state.sections:
        st.warning("Please add at least one section.")
    elif any(not section['title'].strip() for section in st.session_state.sections):
        st.warning("Please provide titles for all sections.")
    elif any(not section['goal'].strip() for section in st.session_state.sections):
        st.warning("Please provide goals for all sections.")
    else:
        try:
            # Get API key from secrets
            client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
            
            processed_sections = []
            progress_bar = st.progress(0, text="Initializing...")

            # Process each section
            for i, section in enumerate(st.session_state.sections):
                progress_text = f"Generating prompt for section: '{section['title']}'..."
                progress_bar.progress((i + 1) / len(st.session_state.sections), text=progress_text)
                
                # Create the special prompt for the AI
                meta_prompt = create_meta_prompt(
                    master_context=master_context,
                    input_files=st.session_state.input_files, # Pass input_files to the prompt
                    section_title=section['title'],
                    section_goal=section['goal'],
                    output_format=section['format']
                )

                # Call OpenAI
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": meta_prompt}],
                    temperature=0.5,
                    response_format={"type": "json_object"} # Force JSON output
                )
                
                # Format the response into the Klarity comment structure
                ai_json_response = response.choices[0].message.content
                comment = format_comment_string(ai_json_response)
                
                processed_sections.append({'title': section['title'], 'comment_string': comment})

            progress_bar.success("All sections processed!")
            
            # Create the Word document in memory
            st.session_state.docx_file = create_docx(
                template_title=st.session_state.sections[0]['title'].split('&')[0].strip(), # Use first part of first section as title
                processed_sections=processed_sections
            )
            st.session_state.file_name = f"Klarity_Template_{st.session_state.sections[0]['title'].split('&')[0].strip().replace(' ', '_')}.docx"

            # Add success message
            st.success("✅ Template generated successfully! Click the download button below to save your file.")

        except openai.AuthenticationError:
            st.error("Authentication Error: Please check your OpenAI API key in the secrets.toml file.")
        except openai.RateLimitError:
            st.error("Rate limit exceeded. Please wait a moment and try again.")
        except Exception as e:
            st.error(f"An error occurred: {e}")
            st.info("💡 Tip: Make sure your OpenAI API key is correctly set in the secrets.toml file.")

# Display download button if file is ready
if 'docx_file' in st.session_state:
    st.download_button(
        label="📥 Download Template.docx",
        data=st.session_state.docx_file,
        file_name=st.session_state.file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )